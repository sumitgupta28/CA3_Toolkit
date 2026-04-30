"""
================================================================================
  generate_test_documents.py
  CA3 Top Sheet — Test Document Generator
================================================================================

PURPOSE
-------
Generates 10 filled student .docx files from the CA3 template for testing the
full workflow (extract_to_excel.py → fill_marks_and_export.py).

Each document simulates a student who has:
  - Filled in all fields in the top information table
  - "Signed" the bottom section with their name and date

REQUIREMENTS
------------
    pip install python-docx

USAGE
-----
    python generate_test_documents.py

OUTPUT
------
    students/
    ├── Student_1083062235.docx   (Aarav Sharma)
    ├── Student_1083062236.docx   (Priya Mehta)
    ├── ...
    └── Student_1083062244.docx   (Megha Singh)

CONFIGURATION
-------------
Edit the sections below to customise:
  - TEMPLATE_PATH  : path to the original blank .docx template
  - OUTPUT_FOLDER  : where the generated files will be saved
  - COMMON_DATA    : shared fields (college, subject, teacher, etc.)
  - STUDENTS       : list of individual student details
================================================================================
"""

import sys
import zipfile
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── dependency check ──────────────────────────────────────────────────────────
try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx is not installed.")
    print("       Run:  pip install python-docx")
    sys.exit(1)


# ══════════════════════════════════════════════════════════════════════════════
#  CONFIGURATION — edit this section to match your setup
# ══════════════════════════════════════════════════════════════════════════════

# Path to the original blank CA3 template .docx
# Place this script in the same folder as the template, or update the path.
TEMPLATE_PATH = "Seal_Top_Sheet_template.docx"

# Folder where generated student documents will be saved
OUTPUT_FOLDER = "students"

# Whether to also create a .zip archive of all documents
CREATE_ZIP = True
ZIP_NAME   = "students_test_documents.zip"

# ── Fields shared by all students ─────────────────────────────────────────────
COMMON_DATA = {
    "college"    : "Model College of Technology, Kolkata - Code: MCT042",
    "program"    : "Bachelor of Computer Applications (BCA)",
    "year_sem"   : "3rd Year / 6th Semester",
    "subject"    : "Data Structures and Algorithms",
    "paper_code" : "BCA601",
    "teacher"    : "Dr. Rajesh Kumar",
    "full_marks" : "30",
    "duration"   : "1.5 Hours",
    "date_exam"  : "15-04-2025",
}

# ── Individual student data ────────────────────────────────────────────────────
# upid     : written into the document template (informational only)
# name     : full name of the student
# roll     : roll number — this is the unique key used to match Excel rows later
# mobile   : contact number
# date_exam: date the student signed/submitted (can differ per student)
STUDENTS = [
    {
        "upid"     : "UPID2025001",
        "name"     : "Aarav Sharma",
        "roll"     : "1083062235",
        "mobile"   : "9876543210",
        "date_exam": "15-04-2025",
    },
    {
        "upid"     : "UPID2025002",
        "name"     : "Priya Mehta",
        "roll"     : "1083062236",
        "mobile"   : "9876543211",
        "date_exam": "15-04-2025",
    },
    {
        "upid"     : "UPID2025003",
        "name"     : "Rohan Das",
        "roll"     : "1083062237",
        "mobile"   : "9876543212",
        "date_exam": "16-04-2025",
    },
    {
        "upid"     : "UPID2025004",
        "name"     : "Sneha Bose",
        "roll"     : "1083062238",
        "mobile"   : "9876543213",
        "date_exam": "16-04-2025",
    },
    {
        "upid"     : "UPID2025005",
        "name"     : "Vikram Yadav",
        "roll"     : "1083062239",
        "mobile"   : "9876543214",
        "date_exam": "15-04-2025",
    },
    {
        "upid"     : "UPID2025006",
        "name"     : "Ananya Roy",
        "roll"     : "1083062240",
        "mobile"   : "9876543215",
        "date_exam": "15-04-2025",
    },
    {
        "upid"     : "UPID2025007",
        "name"     : "Karan Gupta",
        "roll"     : "1083062241",
        "mobile"   : "9876543216",
        "date_exam": "17-04-2025",
    },
    {
        "upid"     : "UPID2025008",
        "name"     : "Deepika Nair",
        "roll"     : "1083062242",
        "mobile"   : "9876543217",
        "date_exam": "17-04-2025",
    },
    {
        "upid"     : "UPID2025009",
        "name"     : "Arjun Patel",
        "roll"     : "1083062243",
        "mobile"   : "9876543218",
        "date_exam": "15-04-2025",
    },
    {
        "upid"     : "UPID2025010",
        "name"     : "Megha Singh",
        "roll"     : "1083062244",
        "mobile"   : "9876543219",
        "date_exam": "16-04-2025",
    },
]


# ══════════════════════════════════════════════════════════════════════════════
#  HELPER FUNCTIONS
# ══════════════════════════════════════════════════════════════════════════════

def get_para_text(para):
    """Return the full plain text of a paragraph by joining all runs."""
    return "".join(run.text for run in para.runs).strip()


def clear_and_set_cell(cell, text, bold=True, font_size=10):
    """
    Clear all existing runs from a table cell's first paragraph
    and write new text with the given formatting.
    """
    para = cell.paragraphs[0]
    for run in para.runs:
        run._r.getparent().remove(run._r)
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)


def clear_and_set_para(para, text, italic=False, font_size=10):
    """
    Clear all existing runs from a paragraph and write new text.
    """
    for run in para.runs:
        run._r.getparent().remove(run._r)
    run = para.add_run(text)
    run.italic = italic
    run.font.size = Pt(font_size)


def insert_signature_paragraph(cell_para, student_name, sign_date):
    """
    Insert a new paragraph with italic blue text directly before the given
    'Signature of the student with date' label paragraph.
    This simulates the student's handwritten signature.
    """
    p_elem  = cell_para._p
    parent  = p_elem.getparent()
    insert_pos = list(parent).index(p_elem)

    # Build <w:p><w:r><w:rPr>...</w:rPr><w:t>...</w:t></w:r></w:p>
    new_p   = OxmlElement("w:p")
    new_r   = OxmlElement("w:r")
    new_rpr = OxmlElement("w:rPr")

    # Italic
    new_rpr.append(OxmlElement("w:i"))

    # Font size (20 half-points = 10pt)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "22")           # 11pt
    new_rpr.append(sz)

    # Dark blue colour to distinguish from printed text
    col = OxmlElement("w:color")
    col.set(qn("w:val"), "1F3864")
    new_rpr.append(col)

    new_r.append(new_rpr)

    # Text: "Name  |  DD-MM-YYYY"
    new_t = OxmlElement("w:t")
    new_t.text = f"{student_name}   |   {sign_date}"
    new_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    new_r.append(new_t)

    new_p.append(new_r)
    parent.insert(insert_pos, new_p)


# ══════════════════════════════════════════════════════════════════════════════
#  CORE: Build one filled document from the template
# ══════════════════════════════════════════════════════════════════════════════

def build_student_document(template_path, student, common):
    """
    Open the template and fill in:
      1. College Code & Name paragraph
      2. Top information table (6 rows × 2 columns)
      3. Student signature in the bottom table

    Returns the filled Document object.
    """
    doc    = Document(template_path)
    tables = doc.tables

    # ── 1. College Code & Name (standalone paragraph above table) ──────────
    for para in doc.paragraphs:
        text = get_para_text(para)
        if "College Code" in text or "College" in text:
            clear_and_set_para(
                para,
                f"College Code & Name: {common['college']}",
                font_size=11,
            )
            break   # only the first matching paragraph

    # ── 2. Top information table ────────────────────────────────────────────
    #
    #  Row 0 | Program Name:           | Year/Semester:
    #  Row 1 | Subject (course):       | Paper (course) Code:
    #  Row 2 | UPID:                   | Date of Examination:
    #  Row 3 | Name of the Student:    | Roll Number:
    #  Row 4 | Subject Teacher:        | Mobile Number:
    #  Row 5 | Full Marks:             | Duration:
    #
    info_table = tables[0]

    clear_and_set_cell(info_table.rows[0].cells[0],
                       f"Program Name: {common['program']}")
    clear_and_set_cell(info_table.rows[0].cells[1],
                       f"Year/Semester: {common['year_sem']}")

    clear_and_set_cell(info_table.rows[1].cells[0],
                       f"Subject (course): {common['subject']}")
    clear_and_set_cell(info_table.rows[1].cells[1],
                       f"Paper (course) Code: {common['paper_code']}")

    clear_and_set_cell(info_table.rows[2].cells[0],
                       f"UPID: {student['upid']}")
    clear_and_set_cell(info_table.rows[2].cells[1],
                       f"Date of Examination: {student['date_exam']}")

    clear_and_set_cell(info_table.rows[3].cells[0],
                       f"Name of the Student: {student['name']}")
    clear_and_set_cell(info_table.rows[3].cells[1],
                       f"Roll Number: {student['roll']}")

    clear_and_set_cell(info_table.rows[4].cells[0],
                       f"Subject Teacher: {common['teacher']}")
    clear_and_set_cell(info_table.rows[4].cells[1],
                       f"Mobile Number: {student['mobile']}")

    clear_and_set_cell(info_table.rows[5].cells[0],
                       f"Full Marks: {common['full_marks']}")
    clear_and_set_cell(info_table.rows[5].cells[1],
                       f"Duration: {common['duration']}")

    # ── 3. Student signature (inside the bottom table) ──────────────────────
    # The bottom table's left cell contains several paragraphs,
    # the last of which is "Signature of the student with date".
    # We insert a new italic paragraph just before that label.
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if "Signature of the student with date" in get_para_text(para):
                        insert_signature_paragraph(
                            para,
                            student_name=student["name"],
                            sign_date=student["date_exam"],
                        )
                        break   # only insert once per document

    return doc


# ══════════════════════════════════════════════════════════════════════════════
#  MAIN
# ══════════════════════════════════════════════════════════════════════════════

def main():
    # ── validate template path ────────────────────────────────────────────
    template = Path(TEMPLATE_PATH)
    if not template.exists():
        print(f"ERROR: Template not found: {TEMPLATE_PATH}")
        print("       Update TEMPLATE_PATH at the top of this script.")
        sys.exit(1)

    # ── create output folder ──────────────────────────────────────────────
    out_dir = Path(OUTPUT_FOLDER)
    out_dir.mkdir(parents=True, exist_ok=True)

    # ── check for duplicate Roll Numbers ─────────────────────────────────
    rolls = [s["roll"] for s in STUDENTS]
    if len(rolls) != len(set(rolls)):
        duplicates = [r for r in rolls if rolls.count(r) > 1]
        print(f"ERROR: Duplicate Roll Numbers found: {set(duplicates)}")
        print("       Each student must have a unique Roll Number.")
        sys.exit(1)

    # ── generate documents ────────────────────────────────────────────────
    print(f"\nTemplate : {template}")
    print(f"Output   : {out_dir.resolve()}")
    print(f"Students : {len(STUDENTS)}")
    print("-" * 56)

    generated = []
    errors    = []

    for student in STUDENTS:
        fname = f"Student_{student['roll']}.docx"
        fpath = out_dir / fname
        try:
            doc = build_student_document(template, student, COMMON_DATA)
            doc.save(str(fpath))
            generated.append(fpath)
            print(f"  ✓  {fname:<35}  {student['name']}")
        except Exception as e:
            errors.append((fname, str(e)))
            print(f"  ✗  {fname:<35}  ERROR: {e}")

    # ── summary ───────────────────────────────────────────────────────────
    print("-" * 56)
    print(f"\n  Generated : {len(generated)} document(s)")
    if errors:
        print(f"  Errors    : {len(errors)}")
        for name, err in errors:
            print(f"             {name} — {err}")

    # ── optional zip ─────────────────────────────────────────────────────
    if CREATE_ZIP and generated:
        zip_path = Path(ZIP_NAME)
        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
            for fpath in generated:
                # Store inside a 'students/' subfolder in the zip
                zf.write(fpath, f"students/{fpath.name}")
        size_kb = zip_path.stat().st_size / 1024
        print(f"\n  Zip created : {zip_path.name}  ({size_kb:.1f} KB)")

    print(f"\n  Done! Files saved to: {out_dir.resolve()}/")
    print()
    print("  Next step: copy the contents of the students/ folder into your")
    print("  project's students/ directory, then run:")
    print("      python extract_to_excel.py")


if __name__ == "__main__":
    main()
