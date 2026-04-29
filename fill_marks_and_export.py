"""
STEPS 4 & 5 — Read filled marks.xlsx → update each student .docx → export PDF
-------------------------------------------------------------------------------
Usage:
    python fill_marks_and_export.py

Before running:
  - Fill in marks.xlsx (yellow columns: Awarded Q1a–Q7, Strengths, Improvements, Measures)
  - Place your signature image as 'examiner_signature.png' in this folder
  - LibreOffice must be installed for PDF export

Output:
  - Updated .docx files in output/docx/
  - PDF files in output/pdf/
"""

import os
import sys
import shutil
import glob
import platform
import subprocess
from pathlib import Path
from copy import deepcopy
from datetime import date

# ── dependency check ─────────────────────────────────────────────────────────
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import lxml.etree as etree
except ImportError:
    print("ERROR: python-docx not found. Run:  pip install python-docx openpyxl")
    sys.exit(1)
try:
    import openpyxl
except ImportError:
    print("ERROR: openpyxl not found. Run:  pip install python-docx openpyxl")
    sys.exit(1)
try:
    from docx2pdf import convert as docx2pdf_convert
    DOCX2PDF_AVAILABLE = True
except ImportError:
    DOCX2PDF_AVAILABLE = False

# ── configuration ─────────────────────────────────────────────────────────────
STUDENTS_FOLDER    = "students"               # original student .docx files
MARKS_EXCEL        = str(Path("CA3_Marks") / "marks.xlsx")  # your filled-in spreadsheet
SIGNATURE_IMAGE    = "examiner_signature.png"               # examiner signature image
OUTPUT_DOCX_FOLDER = Path("output") / "docx"                # updated .docx files
OUTPUT_PDF_FOLDER  = Path("CA3_Marks_Pdf")                  # final PDFs

SIGNATURE_WIDTH_INCHES = 1.6   # signature image width in the document
TODAY_STR = date.today().strftime("%d-%m-%Y")

QUESTION_LABELS = ["1a", "1b", "1c", "1d", "1e", "2", "3", "4", "5", "6", "7"]


# ─────────────────────────────────────────────────────────────────────────────
# Excel helpers
# ─────────────────────────────────────────────────────────────────────────────

def load_marks_excel(excel_path):
    """
    Returns a dict keyed by UPID (uppercase, stripped).
    Each value is a dict of all columns from that row.
    """
    wb = openpyxl.load_workbook(excel_path, data_only=True)
    ws = wb.active

    headers = [str(cell.value).strip() if cell.value else "" for cell in ws[1]]
    records = {}

    for row in ws.iter_rows(min_row=2, values_only=True):
        row_dict = {headers[i]: (str(v).strip() if v is not None else "") for i, v in enumerate(row)}
        upid = row_dict.get("UPID", "").strip().upper()
        if upid:
            records[upid] = row_dict
        elif row_dict.get("upid", "").strip():
            upid = row_dict["upid"].strip().upper()
            records[upid] = row_dict

    return records


# ─────────────────────────────────────────────────────────────────────────────
# DOCX helpers — find text inside paragraphs (handles fragmented runs)
# ─────────────────────────────────────────────────────────────────────────────

def para_full_text(para):
    return "".join(r.text for r in para.runs)


def set_para_text(para, new_text):
    """Replace all runs in a paragraph with a single run containing new_text."""
    # preserve formatting from first run if available
    rpr_xml = None
    if para.runs:
        rpr = para.runs[0]._r.find(qn("w:rPr"))
        if rpr is not None:
            rpr_xml = deepcopy(rpr)
    # clear existing runs
    for run in para.runs:
        run._r.getparent().remove(run._r)
    # add new run
    new_r = OxmlElement("w:r")
    if rpr_xml is not None:
        new_r.append(rpr_xml)
    t = OxmlElement("w:t")
    t.text = new_text
    if new_text and (new_text[0] == " " or new_text[-1] == " "):
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    new_r.append(t)
    para._p.append(new_r)


def get_or_create_run_in_cell(cell):
    """Return the first run in a cell's first paragraph, creating one if needed."""
    para = cell.paragraphs[0]
    if para.runs:
        return para.runs[0]
    run = para.add_run()
    return run


def set_cell_value(cell, value):
    """Set the text value of a table cell cleanly."""
    para = cell.paragraphs[0]
    # remove all runs
    for run in para.runs:
        run._r.getparent().remove(run._r)
    if value:
        run = para.add_run(str(value))
        run.font.size = Pt(10)


# ─────────────────────────────────────────────────────────────────────────────
# Core document update
# ─────────────────────────────────────────────────────────────────────────────

def update_document(src_path, dest_path, row_data, signature_path):
    """
    Open the student docx, fill in:
      - Marks Awarded column in the tabulation table (table index 1)
      - Examiner feedback paragraphs
      - Examiner signature image + date
    Save to dest_path.
    """
    doc = Document(src_path)
    tables = doc.tables

    # ── Marks Tabulation table (table[2]) ────────────────────────────────
    # Table 0 = info, Table 1 = Assessment Rubrics (untouched), Table 2 = marks
    if len(tables) >= 3:
        marks_table = tables[2]
        for i, label in enumerate(QUESTION_LABELS):
            row_i = i + 1   # row 0 is header
            try:
                cells = marks_table.rows[row_i].cells
                col_map = [
                    (2, f"Awarded {label.upper()}"),
                    (3, f"Course Outcome {label.upper()}"),
                    (4, f"Blooms Level {label.upper()}"),
                    (5, f"Remarks {label.upper()}"),
                    (6, f"AR Ref {label.upper()}"),
                ]
                for col_idx, excel_key in col_map:
                    if col_idx < len(cells):
                        val = row_data.get(excel_key, "").strip()
                        if val:
                            set_cell_value(cells[col_idx], val)
            except IndexError:
                pass

    # ── Examiner feedback paragraphs ─────────────────────────────────────
    # We look for the paragraph IMMEDIATELY after the label paragraphs
    feedback_map = {
        "Strengths of the Student:": ("Strengths",              "strengths"),
        "Areas for Improvement:":    ("Areas for Improvement",  "areas_for_improvement"),
        "Suggested Corrective Measures:": ("Corrective Measures", "corrective_measures"),
    }
    # Build lookup: stripped full text → paragraph index
    para_texts = [para_full_text(p).strip() for p in doc.paragraphs]

    for label_text, (_, excel_key) in feedback_map.items():
        # find the label paragraph
        for idx, pt in enumerate(para_texts):
            if label_text.lower() in pt.lower():
                # The next paragraph is the value paragraph
                if idx + 1 < len(doc.paragraphs):
                    val = row_data.get(excel_key, "").strip()
                    # Try both capitalised and original key variants
                    if not val:
                        # try Excel header variants
                        for possible_key in [excel_key,
                                             excel_key.replace("_", " ").title(),
                                             excel_key.replace("_", " ")]:
                            val = row_data.get(possible_key, "").strip()
                            if val:
                                break
                    if val:
                        set_para_text(doc.paragraphs[idx + 1], val)
                break

    # ── Examiner signature & date ─────────────────────────────────────────
    # Find the paragraph containing "Signature of the Examiner with date"
    for idx, pt in enumerate(para_texts):
        if "Signature of the Examiner" in pt:
            sig_para = doc.paragraphs[idx]
            # Clear the paragraph and replace with signature image + date text
            for run in sig_para.runs:
                run._r.getparent().remove(run._r)

            if signature_path and Path(signature_path).exists():
                run_img = sig_para.add_run()
                run_img.add_picture(signature_path, width=Inches(SIGNATURE_WIDTH_INCHES))

            run_date = sig_para.add_run(f"  {TODAY_STR}")
            run_date.font.size = Pt(9)
            break

    doc.save(dest_path)


# ─────────────────────────────────────────────────────────────────────────────
# PDF conversion
# ─────────────────────────────────────────────────────────────────────────────

def find_libreoffice():
    """Return the LibreOffice executable path for Mac or Windows."""
    system = platform.system()
    candidates = []

    if system == "Darwin":   # macOS
        candidates = [
            "/Applications/LibreOffice.app/Contents/MacOS/soffice",
            "/usr/local/bin/soffice",
            "soffice",
        ]
    elif system == "Windows":
        program_files = [
            os.environ.get("PROGRAMFILES", "C:\\Program Files"),
            os.environ.get("PROGRAMFILES(X86)", "C:\\Program Files (x86)"),
        ]
        for pf in program_files:
            candidates.append(os.path.join(pf, "LibreOffice", "program", "soffice.exe"))
        candidates.append("soffice.exe")
    else:  # Linux / other
        candidates = ["libreoffice", "soffice"]

    for c in candidates:
        if shutil.which(c) or (os.path.isfile(c)):
            return c
    return None


def convert_to_pdf(docx_path, pdf_output_dir):
    """Convert a .docx to PDF — tries LibreOffice first, then docx2pdf."""
    lo = find_libreoffice()
    if lo:
        cmd = [
            lo,
            "--headless",
            "--convert-to", "pdf",
            "--outdir", str(pdf_output_dir),
            str(docx_path),
        ]
        try:
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
            if result.returncode == 0:
                return True, ""
            return False, result.stderr.strip()
        except subprocess.TimeoutExpired:
            return False, "Timeout"
        except Exception as e:
            return False, str(e)

    if DOCX2PDF_AVAILABLE:
        try:
            pdf_path = Path(pdf_output_dir) / (Path(docx_path).stem + ".pdf")
            docx2pdf_convert(str(docx_path), str(pdf_path))
            return True, ""
        except Exception as e:
            return False, str(e)

    return False, "No PDF converter available (install LibreOffice or docx2pdf)"


# ─────────────────────────────────────────────────────────────────────────────
# Main
# ─────────────────────────────────────────────────────────────────────────────

def main():
    # ── pre-flight checks ────────────────────────────────────────────────
    students_folder = Path(STUDENTS_FOLDER)
    if not students_folder.exists():
        print(f"ERROR: Folder '{STUDENTS_FOLDER}/' not found.")
        sys.exit(1)

    if not Path(MARKS_EXCEL).exists():
        print(f"ERROR: '{MARKS_EXCEL}' not found. Ensure marks.xlsx is in the CA3_Marks/ folder.")
        sys.exit(1)

    sig_path = SIGNATURE_IMAGE if Path(SIGNATURE_IMAGE).exists() else None
    if not sig_path:
        print(f"⚠  WARNING: '{SIGNATURE_IMAGE}' not found — signature image will be skipped.")
        print("            Place your signature PNG in the same folder as this script.\n")

    lo = find_libreoffice()
    if not lo and not DOCX2PDF_AVAILABLE:
        print("⚠  WARNING: No PDF converter found — PDFs will NOT be generated.")
        print("            Install docx2pdf:  pip install docx2pdf\n")
    elif not lo and DOCX2PDF_AVAILABLE:
        print("ℹ  LibreOffice not found — using docx2pdf for PDF conversion.\n")

    # ── create output dirs ────────────────────────────────────────────────
    OUTPUT_DOCX_FOLDER.mkdir(parents=True, exist_ok=True)
    OUTPUT_PDF_FOLDER.mkdir(parents=True, exist_ok=True)

    # ── load Excel data ───────────────────────────────────────────────────
    print(f"\n📊  Loading marks from '{MARKS_EXCEL}' ...")
    marks_data = load_marks_excel(MARKS_EXCEL)
    print(f"    {len(marks_data)} student record(s) loaded.")

    if not marks_data:
        print("ERROR: No data found in Excel. Make sure UPID column is filled.")
        sys.exit(1)

    # ── find all student docx files ───────────────────────────────────────
    docx_files = sorted(glob.glob(str(students_folder / "*.docx")))
    if not docx_files:
        print(f"ERROR: No .docx files in '{STUDENTS_FOLDER}/'")
        sys.exit(1)

    print(f"📂  Found {len(docx_files)} .docx file(s) in '{STUDENTS_FOLDER}/'")
    print("-" * 60)

    ok_count = 0
    skip_count = 0
    pdf_ok = 0
    pdf_fail = 0

    for fp in docx_files:
        fname = Path(fp).name
        print(f"\n  File: {fname}")

        # read UPID from document
        try:
            doc_tmp = Document(fp)
            tables  = doc_tmp.tables
            raw_upid = ""
            if tables:
                # UPID is row 2, col 0 of the first table
                raw_upid = tables[0].rows[2].cells[0].text.strip()
                # strip the label
                if ":" in raw_upid:
                    raw_upid = raw_upid.split(":", 1)[1].strip()
            upid = raw_upid.upper()
        except Exception as e:
            print(f"    ⚠  Could not read UPID: {e}  — skipping.")
            skip_count += 1
            continue

        if not upid:
            print(f"    ⚠  UPID is blank — skipping.")
            skip_count += 1
            continue

        row_data = marks_data.get(upid)
        if not row_data:
            print(f"    ⚠  UPID '{upid}' not found in Excel — skipping.")
            skip_count += 1
            continue

        print(f"    UPID={upid}  →  matched in Excel ✓")

        # ── update docx ──────────────────────────────────────────────────
        out_docx = OUTPUT_DOCX_FOLDER / fname
        try:
            update_document(fp, out_docx, row_data, sig_path)
            print(f"    📄  Updated docx saved: {out_docx}")
            ok_count += 1
        except Exception as e:
            print(f"    ✗  Failed to update docx: {e}")
            skip_count += 1
            continue

        # ── convert to PDF ───────────────────────────────────────────────
        if lo or DOCX2PDF_AVAILABLE:
            success, err = convert_to_pdf(out_docx, OUTPUT_PDF_FOLDER)
            if success:
                pdf_name = Path(fname).stem + ".pdf"
                print(f"    📑  PDF exported:  {OUTPUT_PDF_FOLDER / pdf_name}")
                pdf_ok += 1
            else:
                print(f"    ⚠  PDF conversion failed: {err}")
                pdf_fail += 1
        else:
            print(f"    ⏭  PDF skipped (no PDF converter available)")

    # ── summary ───────────────────────────────────────────────────────────
    print("\n" + "=" * 60)
    print(f"✅  Done!  {ok_count} document(s) updated successfully.")
    if skip_count:
        print(f"⚠   {skip_count} file(s) skipped (check warnings above).")
    if lo or DOCX2PDF_AVAILABLE:
        print(f"📑  PDFs generated: {pdf_ok}   |   Failed: {pdf_fail}")
        if pdf_ok:
            print(f"    PDFs saved to: {OUTPUT_PDF_FOLDER}/")
    print(f"    Updated .docx files: {OUTPUT_DOCX_FOLDER}/")
    print("=" * 60)


if __name__ == "__main__":
    main()
