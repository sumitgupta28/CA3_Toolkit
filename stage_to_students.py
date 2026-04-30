"""
stage_to_students.py

Scans stage_students_files/ for .docx files, detects their true format,
converts non-docx files to proper .docx, and moves all files into students/.

Handled formats:
  - Proper .docx (ZIP/OOXML)       → moved directly (no external tools)
  - .dotx template (ZIP/OOXML)     → content-type fixed in-memory (no external tools)
  - OLE2 binary .doc               → converted via LibreOffice headless

Requires LibreOffice for OLE2 binary conversion (free, no Microsoft Office needed):
  macOS:   https://www.libreoffice.org/download/
  Windows: https://www.libreoffice.org/download/
  Linux:   sudo apt install libreoffice
"""

import os
import platform
import shutil
import subprocess
import tempfile
import zipfile
from pathlib import Path

try:
    from docx import Document as _DocxDocument
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

STAGE_FOLDER = "stage_students_files"
STUDENTS_FOLDER = "students"

OLE2_MAGIC = b"\xD0\xCF\x11\xE0\xA1\xB1\x1A\xE1"


# ---------------------------------------------------------------------------
# LibreOffice detection — compatible with macOS, Windows, and Linux
# ---------------------------------------------------------------------------

def find_libreoffice():
    system = platform.system()
    if system == "Darwin":
        candidates = [
            "/Applications/LibreOffice.app/Contents/MacOS/soffice",
            "/usr/local/bin/soffice",
            "soffice",
        ]
    elif system == "Windows":
        program_files = [
            os.environ.get("PROGRAMFILES", "C:\\Program Files"),
            os.environ.get("PROGRAMFILES(X86)", "C:\\Program Files (x86)"),
        ]
        candidates = [os.path.join(pf, "LibreOffice", "program", "soffice.exe") for pf in program_files]
        candidates.append("soffice.exe")
    else:
        candidates = ["libreoffice", "soffice"]

    for c in candidates:
        if shutil.which(c) or os.path.isfile(c):
            return c
    return None


# ---------------------------------------------------------------------------
# Roll Number validation
# ---------------------------------------------------------------------------

def get_roll_number(path: Path) -> str:
    """Read Roll Number from row 3, col 1 of the top info table."""
    if not _DOCX_AVAILABLE:
        return "unknown"
    try:
        doc = _DocxDocument(str(path))
        if doc.tables:
            raw = doc.tables[0].rows[3].cells[1].text.strip()
            if ":" in raw:
                return raw.split(":", 1)[1].strip()
            return raw
    except Exception:
        pass
    return ""


# ---------------------------------------------------------------------------
# Format detection
# ---------------------------------------------------------------------------

def detect_format(path: Path) -> str:
    with open(path, "rb") as f:
        magic = f.read(8)

    if magic[:2] == b"PK":
        try:
            with zipfile.ZipFile(path) as z:
                ct = z.read("[Content_Types].xml").decode("utf-8", errors="replace")
            if "wordprocessingml.document.main" in ct:
                return "docx"
            if "wordprocessingml.template.main" in ct:
                return "dotx"
            return "other_ooxml"
        except Exception:
            return "unknown"

    if magic[:8] == OLE2_MAGIC:
        return "doc_ole2"

    return "unknown"


# ---------------------------------------------------------------------------
# Converters
# ---------------------------------------------------------------------------

def convert_dotx_to_docx(src: Path, dst: Path):
    with zipfile.ZipFile(src, "r") as zin, zipfile.ZipFile(dst, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "[Content_Types].xml":
                data = data.replace(
                    b"wordprocessingml.template.main+xml",
                    b"wordprocessingml.document.main+xml",
                )
            zout.writestr(item, data)


def convert_doc_to_docx(src: Path, dst: Path, lo: str) -> tuple[bool, str]:
    with tempfile.TemporaryDirectory() as tmp:
        result = subprocess.run(
            [lo, "--headless", "--convert-to", "docx", str(src.resolve()), "--outdir", tmp],
            capture_output=True,
            text=True,
        )
        converted = Path(tmp) / (src.stem + ".docx")
        if result.returncode != 0 or not converted.exists():
            err = result.stderr.strip() or result.stdout.strip() or "unknown error"
            return False, f"LibreOffice conversion failed: {err}"
        shutil.move(str(converted), str(dst))
    return True, "OLE2 binary → .docx via LibreOffice"


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    stage_dir = Path(STAGE_FOLDER)
    students_dir = Path(STUDENTS_FOLDER)

    if not stage_dir.exists():
        print(f"ERROR: Folder '{STAGE_FOLDER}' not found.")
        return

    students_dir.mkdir(exist_ok=True)

    files = sorted(stage_dir.glob("*.docx"))
    if not files:
        print(f"No .docx files found in '{STAGE_FOLDER}'.")
        return

    # Pre-check: locate LibreOffice only if OLE2 files are present
    needs_lo = any(detect_format(f) == "doc_ole2" for f in files)
    lo = find_libreoffice() if needs_lo else None
    if needs_lo and not lo:
        print("ERROR: LibreOffice is required to convert OLE2 binary .doc files.")
        print("       Download: https://www.libreoffice.org/download/")
        return

    print(f"\nProcessing {len(files)} files from {STAGE_FOLDER}/ → {STUDENTS_FOLDER}/\n")

    moved = converted = failed = skipped = 0
    skipped_files = []

    for path in files:
        fmt = detect_format(path)
        dst = students_dir / path.name

        if fmt in ("docx", "other_ooxml"):
            shutil.copy2(str(path), str(dst))
            placed = True
        elif fmt == "dotx":
            convert_dotx_to_docx(path, dst)
            placed = True
        elif fmt == "doc_ole2":
            ok, err_msg = convert_doc_to_docx(path, dst, lo)
            if ok:
                placed = True
            else:
                print(f"  [FAILED]          {path.name}  {err_msg}")
                failed += 1
                continue
        else:
            print(f"  [SKIPPED]         {path.name}  (unrecognised format: {fmt})")
            failed += 1
            continue

        if placed:
            roll = get_roll_number(dst)
            if not roll:
                dst.unlink()
                reason = "Roll Number is blank in the document"
                print(f"  [SKIPPED]         {path.name}  ({reason})")
                skipped_files.append((path.name, reason))
                skipped += 1
            elif fmt in ("docx", "other_ooxml"):
                print(f"  [OK - copied]     {path.name}  (already .docx,  Roll={roll})")
                moved += 1
            else:
                label = ".dotx template → .docx" if fmt == "dotx" else "OLE2 binary → .docx via LibreOffice"
                print(f"  [OK - converted]  {path.name}  ({label},  Roll={roll})")
                converted += 1

    print(f"\nSummary: {moved} copied, {converted} converted, {skipped} skipped, {failed} failed")
    if skipped_files:
        print("\nSkipped files:")
        for name, reason in skipped_files:
            print(f"  • {name} — {reason}")


if __name__ == "__main__":
    main()
